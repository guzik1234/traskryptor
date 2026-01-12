"""
Moduł do transkrypcji plików audio z zapisem do DOCX i PDF
"""
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import threading
import os
import numpy as np
import wave
import librosa
from transcription_model import TranscriptionModel
from docx import Document

# Import opcjonalny - jeśli nie ma docx2pdf, będziemy tylko zapisywać DOCX
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("Ostrzeżenie: docx2pdf niedostępny. Tylko DOCX będzie dostępny.")


class AudioFileTranscriptionGUI:
    """GUI do transkrypcji plików audio"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Transkrypcja z pliku audio")
        self.root.geometry("600x600")
        self.root.resizable(False, False)
        
        self.audio_file = None
        self.language = None
        self.transcription_model = None
        self.is_processing = False
        
        self._create_widgets()
    
    def _create_widgets(self):
        """Tworzy widgety GUI"""
        
        # Tytuł
        title = tk.Label(
            self.root, 
            text=" Transkrypcja z pliku audio",
            font=("Arial", 18, "bold")
        )
        title.pack(pady=20)
        
        # Instrukcja
        info = tk.Label(
            self.root,
            text="Wybierz plik audio i język, a następnie rozpocznij transkrypcję.\n"
                 "Wynik zostanie zapisany jako DOCX i PDF.",
            font=("Arial", 10),
            fg="gray"
        )
        info.pack(pady=5)
        
        # Ramka wyboru pliku
        file_frame = tk.LabelFrame(self.root, text="Plik audio", font=("Arial", 11, "bold"))
        file_frame.pack(pady=15, padx=20, fill="x")
        
        self.file_label = tk.Label(
            file_frame, 
            text="Nie wybrano pliku",
            font=("Arial", 10),
            fg="gray"
        )
        self.file_label.pack(pady=10)
        
        select_file_btn = tk.Button(
            file_frame,
            text=" Wybierz plik audio",
            command=self.select_audio_file,
            font=("Arial", 11),
            bg="#2196F3",
            fg="white",
            cursor="hand2"
        )
        select_file_btn.pack(pady=10)
        
        # Ramka wyboru języka
        lang_frame = tk.LabelFrame(self.root, text="Język transkrypcji", font=("Arial", 11, "bold"))
        lang_frame.pack(pady=15, padx=20, fill="x")
        
        self.language_var = tk.StringVar(value="polski")
        
        polski_radio = tk.Radiobutton(
            lang_frame,
            text=" Polski",
            variable=self.language_var,
            value="polski",
            font=("Arial", 11),
            command=self.on_language_change
        )
        polski_radio.pack(pady=5)
        
        angielski_radio = tk.Radiobutton(
            lang_frame,
            text=" Angielski",
            variable=self.language_var,
            value="angielski",
            font=("Arial", 11),
            command=self.on_language_change
        )
        angielski_radio.pack(pady=5)
        
        # Przycisk rozpoczęcia
        self.start_btn = tk.Button(
            self.root,
            text=" Rozpocznij transkrypcję",
            command=self.start_transcription,
            font=("Arial", 13, "bold"),
            bg="#4CAF50",
            fg="white",
            width=25,
            height=2,
            cursor="hand2",
            state="disabled"
        )
        self.start_btn.pack(pady=15)
        
        # Pasek postępu
        self.progress = ttk.Progressbar(
            self.root,
            mode='indeterminate',
            length=400
        )
        self.progress.pack(pady=10)
        
        # Status
        self.status_label = tk.Label(
            self.root,
            text="",
            font=("Arial", 9),
            fg="blue"
        )
        self.status_label.pack(pady=5)
        
        # Przycisk powrotu
        back_btn = tk.Button(
            self.root,
            text=" Powrót do menu",
            command=self.back_to_menu,
            font=("Arial", 10),
            bg="#f44336",
            fg="white"
        )
        back_btn.pack(pady=10)
    
    def on_language_change(self):
        """Obsługuje zmianę języka"""
        pass
    
    def select_audio_file(self):
        """Wybiera plik audio"""
        file_path = filedialog.askopenfilename(
            title="Wybierz plik audio",
            filetypes=[
                ("Pliki audio", "*.wav *.mp3 *.m4a *.flac *.ogg"),
                ("Wszystkie pliki", "*.*")
            ]
        )
        
        if file_path:
            self.audio_file = file_path
            filename = os.path.basename(file_path)
            self.file_label.config(text=filename, fg="black")
            self.start_btn.config(state="normal")
            self.status_label.config(text="")
    
    def start_transcription(self):
        """Rozpoczyna transkrypcję w osobnym wątku"""
        if self.is_processing:
            return
        
        if not self.audio_file:
            messagebox.showwarning("Uwaga", "Najpierw wybierz plik audio!")
            return
        
        self.is_processing = True
        self.start_btn.config(state="disabled")
        self.progress.start()
        self.status_label.config(text="Ładowanie modelu...", fg="blue")
        
        # Uruchom w osobnym wątku
        thread = threading.Thread(target=self.process_transcription)
        thread.daemon = True
        thread.start()
    
    def process_transcription(self):
        """Przetwarza transkrypcję"""
        try:
            # Pobierz ustawienia
            self.language = self.language_var.get()
            
            # Załaduj model (używamy domyślnego modelu jak w transkrypcji na żywo)
            self.root.after(0, lambda: self.status_label.config(text="Ładowanie modelu..."))
            self.transcription_model = TranscriptionModel(self.language, "1")
            
            # Wczytaj audio
            self.root.after(0, lambda: self.status_label.config(text="Wczytywanie pliku audio..."))
            audio_data, sample_rate = self.load_audio_file()
            
            # Transkrybuj
            self.root.after(0, lambda: self.status_label.config(text="Transkrybuję audio..."))
            transcription = self.transcription_model.transcribe(audio_data, sample_rate)
            
            # Zapisz wyniki
            self.root.after(0, lambda: self.status_label.config(text="Zapisywanie wyników..."))
            self.save_results(transcription)
            
            # Sukces
            self.root.after(0, lambda: self.on_success())
            
        except Exception as e:
            self.root.after(0, lambda: self.on_error(str(e)))
        finally:
            self.is_processing = False
            self.root.after(0, lambda: self.progress.stop())
            self.root.after(0, lambda: self.start_btn.config(state="normal"))
    
    def load_audio_file(self):
        """Wczytuje plik audio i konwertuje do formatu 16kHz mono"""
        # Użyj librosa do wczytania audio (obsługuje wiele formatów)
        audio_data, sample_rate = librosa.load(self.audio_file, sr=16000, mono=True)
        return audio_data.astype(np.float32), sample_rate
    
    def save_results(self, transcription):
        """Zapisuje transkrypcję do DOCX i PDF"""
        # Przygotuj nazwę pliku wyjściowego
        base_name = os.path.splitext(self.audio_file)[0]
        docx_path = f"{base_name}_transkrypcja.docx"
        pdf_path = f"{base_name}_transkrypcja.pdf"
        
        # Zapisz do DOCX
        doc = Document()
        doc.add_heading('Transkrypcja audio', 0)
        doc.add_paragraph(f"Plik źródłowy: {os.path.basename(self.audio_file)}")
        doc.add_paragraph(f"Język: {self.language}")
        doc.add_paragraph("")
        doc.add_heading('Treść transkrypcji:', 1)
        doc.add_paragraph(transcription)
        doc.save(docx_path)
        
        self.output_docx = docx_path
        
        # Konwertuj do PDF tylko jeśli docx2pdf jest dostępny
        if DOCX2PDF_AVAILABLE:
            try:
                convert(docx_path, pdf_path)
                self.output_pdf = pdf_path
            except Exception as e:
                print(f"Ostrzeżenie: Nie udało się utworzyć PDF: {e}")
                # Kontynuuj mimo błędu - przynajmniej DOCX zostanie zapisany
        else:
            print("Pominięcie konwersji do PDF - docx2pdf niedostępny")
    
    def on_success(self):
        """Wywoływane po udanej transkrypcji"""
        self.progress.stop()
        self.status_label.config(
            text=" Transkrypcja zakończona!",
            fg="green"
        )
        
        msg = f"Transkrypcja została zapisana:\n\n"
        msg += f" {self.output_docx}\n"
        if hasattr(self, 'output_pdf') and os.path.exists(self.output_pdf):
            msg += f" {self.output_pdf}"
        
        messagebox.showinfo("Sukces", msg)
    
    def on_error(self, error_msg):
        """Wywoływane przy błędzie"""
        self.progress.stop()
        self.status_label.config(
            text=" Błąd transkrypcji",
            fg="red"
        )
        messagebox.showerror("Błąd", f"Wystąpił błąd:\n{error_msg}")
    
    def back_to_menu(self):
        """Powrót do głównego menu"""
        self.root.destroy()
        from main import MainApp
        app = MainApp()
        app.run()
    
    def run(self):
        """Uruchamia GUI"""
        self.root.mainloop()


if __name__ == "__main__":
    app = AudioFileTranscriptionGUI()
    app.run()
